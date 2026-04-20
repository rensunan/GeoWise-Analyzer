"""
app.py - 岩土报告智能解析系统
基于大语言模型的表格解析 - 主表/备注/附表识别
"""

import os
import re
import json
import uuid
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Any
from collections import defaultdict
import threading
import warnings
import hashlib

warnings.filterwarnings('ignore')

from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from openai import OpenAI

app = Flask(__name__)
CORS(app)

app.config['UPLOAD_FOLDER'] = './uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.config['SECRET_KEY'] = 'geotechnical-parser-secret-key'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs('./results', exist_ok=True)

# DeepSeek API配置
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-db068eaf4d794e33a0d452203e4d8e9a")
DEEPSEEK_BASE_URL = "https://api.deepseek.com"

# 初始化OpenAI客户端
client = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url=DEEPSEEK_BASE_URL
)

# ============================================================================
# 文档解析器
# ============================================================================

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装")

try:
    import fitz
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("警告: PyMuPDF未安装")


class DocumentParser:
    def __init__(self):
        self.content_keywords = [
            '黏聚力', '内摩擦角', '压缩模量', '孔隙比', '含水量', '液限', '塑限',
            '颗粒级配', '粒径', '筛分', '剪切', '压缩', '载荷', '承载力'
        ]
        self.skip_keywords = ['目录', 'Contents', '图', 'Fig', '表目录', '图目录', '参考文献']

    def parse_document(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        file_ext = Path(file_path).suffix.lower()
        if file_ext == '.docx':
            return self._parse_docx(file_path)
        elif file_ext == '.pdf':
            return self._parse_pdf(file_path)
        else:
            return self._parse_txt(file_path)

    def _parse_docx(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        if not DOCX_AVAILABLE:
            raise ImportError("请安装 python-docx")

        doc = Document(file_path)
        text_paragraphs = []
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if len(text) < 15:
                continue
            if any(kw in text for kw in self.skip_keywords) and len(text) < 50:
                continue
            has_keyword = any(kw in text for kw in self.content_keywords)
            is_table_title = "表" in text and re.search(r'\d+\.\d+', text) is not None
            if has_keyword or is_table_title:
                text_paragraphs.append({
                    "id": f"P_{i+1:04d}",
                    "content": text,
                    "page": 1,
                    "is_table_title": is_table_title
                })

        return text_paragraphs, tables, 1

    def _parse_pdf(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        if not PDF_AVAILABLE:
            raise ImportError("请安装 PyMuPDF")

        doc = fitz.open(file_path)
        text_paragraphs = []
        tables = []
        para_counter = 0

        for page_num, page in enumerate(doc):
            text = page.get_text()
            for para in text.split('\n\n'):
                para = para.strip()
                if len(para) < 20:
                    continue
                if any(kw in para for kw in self.content_keywords):
                    para_counter += 1
                    text_paragraphs.append({
                        "id": f"P_{page_num+1:04d}_{para_counter:03d}",
                        "content": para,
                        "page": page_num + 1,
                        "is_table_title": False
                    })

            page_tables = page.find_tables()
            for tab in page_tables:
                table_data = []
                for row in tab.extract():
                    row_data = [str(cell).strip() if cell else "" for cell in row]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

        doc.close()
        return text_paragraphs, tables, len(doc)

    def _parse_txt(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        text_paragraphs = []
        tables = []

        for i, para in enumerate(content.split('\n\n')):
            para = para.strip()
            if len(para) > 20 and any(kw in para for kw in self.content_keywords):
                text_paragraphs.append({
                    "id": f"P_{i+1:04d}",
                    "content": para,
                    "page": 1,
                    "is_table_title": False
                })

        return text_paragraphs, tables, 1


# ============================================================================
# 表格处理器
# ============================================================================

class TableProcessor:
    """表格处理器"""

    @staticmethod
    def merge_adjacent_rows(table_data: List[List[str]]) -> List[List[str]]:
        """合并完全相同的相邻行，删除多余行"""
        if not table_data:
            return table_data

        result = []
        prev_row = None

        for row in table_data:
            if prev_row is None or row != prev_row:
                result.append(row)
                prev_row = row

        return result

    @staticmethod
    def merge_adjacent_cols(table_data: List[List[str]]) -> List[List[str]]:
        """合并完全相同的相邻列，删除多余列"""
        if not table_data:
            return table_data

        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        padded = []
        for row in table_data:
            while len(row) < cols:
                row.append("")
            padded.append(row)

        cols_to_delete = [False] * cols

        for j in range(cols - 1):
            is_same = True
            for i in range(rows):
                if padded[i][j] != padded[i][j + 1]:
                    is_same = False
                    break
            if is_same:
                cols_to_delete[j + 1] = True

        new_table = []
        for i in range(rows):
            new_row = []
            for j in range(cols):
                if not cols_to_delete[j]:
                    new_row.append(padded[i][j])
            new_table.append(new_row)

        return new_table

    @staticmethod
    def split_by_header(table_data: List[List[str]]) -> List[List[List[str]]]:
        """按相同表头切分表格 - 相似度超过2/3即可"""
        if not table_data or len(table_data) < 2:
            return [table_data]

        header_row = table_data[0]
        header_non_empty = [(j, header_row[j]) for j in range(len(header_row)) if header_row[j]]

        if not header_non_empty:
            return [table_data]

        split_indices = [0]
        for i in range(1, len(table_data)):
            current_row = table_data[i]

            match_count = 0
            for j, header_text in header_non_empty:
                if j < len(current_row) and current_row[j] == header_text:
                    match_count += 1

            similarity = match_count / len(header_non_empty) if header_non_empty else 0

            if similarity >= 2/3:
                split_indices.append(i)

        result = []
        for idx, start in enumerate(split_indices):
            end = split_indices[idx + 1] if idx + 1 < len(split_indices) else len(table_data)
            result.append(table_data[start:end])

        return result


# ============================================================================
# 大语言模型表格解析器
# ============================================================================

class LLMTableParser:
    """基于大语言模型的表格解析器"""

    def __init__(self):
        self.cache = {}  # 缓存相同结构表格的解析结果

    def _get_table_signature(self, table_data: List[List[str]]) -> str:
        """获取表格结构签名（用于判断是否为相同类型表格）"""
        if not table_data:
            return ""

        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        # 提取表头行
        header_row = table_data[0] if rows > 0 else []

        # 签名 = 行数范围 + 列数 + 表头内容哈希
        row_range = "small" if rows < 20 else "medium" if rows < 50 else "large"
        header_hash = hashlib.md5("|".join(header_row).encode()).hexdigest()[:8]

        return f"{row_range}_{cols}_{header_hash}"

    def _call_deepseek(self, table_data: List[List[str]]) -> Dict:
        """调用DeepSeek API解析表格 - 使用OpenAI SDK"""

        # 将表格转换为文本表示
        table_text = []
        for i, row in enumerate(table_data):
            row_text = " | ".join([str(cell) if cell else "" for cell in row])
            table_text.append(f"行{i}: {row_text}")
        table_str = "\n".join(table_text)

        prompt = f"""你是一个岩土工程表格解析专家。请分析以下表格，识别出表格中的"主表"、"备注"和"附表"。

解析规则：
1. 主表：表格的主要内容，通常是包含表头和数据行的核心表格。主表应该有明确的列名和数据行。
2. 备注：表格中包含的键值对信息（如"地下水深度：未见"），或对整个表格起作用的说明文字。
3. 附表：当一个主表无法完全解析表格时，需要拆分成多个子表的情况。

请返回JSON格式结果，格式如下：
{{
    "has_main_table": true/false,
    "main_table": {{
        "start_row": 起始行号,
        "end_row": 结束行号,
        "start_col": 起始列号,
        "end_col": 结束列号,
        "description": "表格描述"
    }},
    "has_remarks": true/false,
    "remarks": [
        {{
            "start_row": 起始行号,
            "end_row": 结束行号,
            "start_col": 起始列号,
            "end_col": 结束列号,
            "content": "备注内容"
        }}
    ],
    "has_sub_tables": true/false,
    "sub_tables": [
        {{
            "start_row": 起始行号,
            "end_row": 结束行号,
            "start_col": 起始列号,
            "end_col": 结束列号,
            "description": "子表描述"
        }}
    ]
}}

表格数据：
{table_str}

注意：只返回JSON，不要有其他内容。行号和列号从0开始计数。"""

        print(f"  [DeepSeek请求] 表格大小: {len(table_data)}行 x {len(table_data[0]) if table_data else 0}列")
        print(f"  [DeepSeek请求] Prompt长度: {len(prompt)}字符")

        try:
            print(f"  [DeepSeek调用] 发送请求到 {DEEPSEEK_BASE_URL}")

            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "你是一个岩土工程表格解析专家，擅长识别表格结构。只返回JSON格式结果。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000,
                stream=False
            )

            print(f"  [DeepSeek响应] 成功")
            content = response.choices[0].message.content
            print(f"  [DeepSeek返回] 原始内容: {content[:500]}...")

            # 提取JSON
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                parsed = json.loads(json_match.group())
                print(f"  [DeepSeek解析] 成功: has_main_table={parsed.get('has_main_table')}, has_remarks={parsed.get('has_remarks')}, has_sub_tables={parsed.get('has_sub_tables')}")
                return parsed
            else:
                print(f"  [DeepSeek错误] 无法从响应中提取JSON")
                return self._get_default_parse_result(table_data)

        except Exception as e:
            print(f"  [DeepSeek异常] {e}")
            import traceback
            traceback.print_exc()
            return self._get_default_parse_result(table_data)

    def _get_default_parse_result(self, table_data: List[List[str]]) -> Dict:
        """默认解析结果（API失败时的降级方案）"""
        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        return {
            "has_main_table": True,
            "main_table": {
                "start_row": 0,
                "end_row": rows - 1,
                "start_col": 0,
                "end_col": cols - 1,
                "description": "整个表格"
            },
            "has_remarks": False,
            "remarks": [],
            "has_sub_tables": False,
            "sub_tables": []
        }

    def parse(self, table_data: List[List[str]]) -> Dict:
        """解析表格，返回主表/备注/附表的位置信息"""
        if not table_data:
            return {
                "has_main_table": False,
                "main_table": None,
                "has_remarks": False,
                "remarks": [],
                "has_sub_tables": False,
                "sub_tables": []
            }

        # 获取表格签名
        signature = self._get_table_signature(table_data)

        # 检查缓存
        if signature in self.cache:
            print(f"  [缓存命中] 使用相同结构表格的解析结果")
            return self.cache[signature]

        # 调用大模型
        print(f"  [调用DeepSeek] 解析表格 {len(table_data)}行 x {len(table_data[0]) if table_data else 0}列")
        result = self._call_deepseek(table_data)

        # 缓存结果
        self.cache[signature] = result

        return result


# ============================================================================
# 表格提取器 - 根据LLM解析结果提取子表格
# ============================================================================

class TableExtractor:
    """根据LLM解析结果提取子表格"""

    @staticmethod
    def extract_sub_table(table_data: List[List[str]], region: Dict) -> List[List[str]]:
        """提取指定区域的子表格"""
        if not region:
            return []

        start_row = region.get("start_row", 0)
        end_row = region.get("end_row", len(table_data) - 1)
        start_col = region.get("start_col", 0)
        end_col = region.get("end_col", 0)

        # 确保范围有效
        start_row = max(0, min(start_row, len(table_data) - 1))
        end_row = max(start_row, min(end_row, len(table_data) - 1))
        start_col = max(0, start_col)
        end_col = max(start_col, end_col)

        result = []
        for i in range(start_row, end_row + 1):
            if i < len(table_data):
                row = table_data[i]
                if start_col < len(row):
                    result.append(row[start_col:min(end_col + 1, len(row))])
                else:
                    result.append([])

        return result


# ============================================================================
# 主解析系统
# ============================================================================

class GeotechnicalParserSystem:
    def __init__(self):
        print("=" * 60)
        print("岩土报告智能解析系统")
        print("基于大语言模型的表格解析 - 主表/备注/附表识别")
        print("=" * 60)

        print("\n[1] 初始化文档解析器...")
        self.doc_parser = DocumentParser()

        print("\n[2] 初始化表格处理器...")
        self.table_processor = TableProcessor()

        print("\n[3] 初始化大语言模型解析器...")
        self.llm_parser = LLMTableParser()

        print("\n[4] 初始化表格提取器...")
        self.table_extractor = TableExtractor()

        print("\n✓ 系统初始化完成！")

    def parse_document(self, file_path: str) -> Dict:
        print(f"\n[解析文档] {file_path}")

        print("  步骤1: 提取文本段落和表格...")
        text_paragraphs, raw_tables, total_pages = self.doc_parser.parse_document(file_path)
        print(f"    文本段落: {len(text_paragraphs)}")
        print(f"    原始表格: {len(raw_tables)}")

        print("  步骤2: 按表头切分表格...")
        print("  步骤3: 对每个子表格合并相邻相同行列...")
        print("  步骤4: 大语言模型解析表格结构...")

        parsed_tables = []
        for idx, table_data in enumerate(raw_tables):
            print(f"    处理表格 T_{idx+1:04d}: 原始 {len(table_data)}行")

            # 第一步：按表头切分
            split_tables = self.table_processor.split_by_header(table_data)
            print(f"      切分为 {len(split_tables)} 个独立表格")

            for sub_idx, sub_table in enumerate(split_tables):
                # 第二步：合并相邻相同行和列
                merged_rows = self.table_processor.merge_adjacent_rows(sub_table)
                merged_table = self.table_processor.merge_adjacent_cols(merged_rows)

                # 第三步：大语言模型解析表格结构
                parse_result = self.llm_parser.parse(merged_table)

                # 第四步：根据解析结果提取子表格
                main_table_data = []
                remark_data = []
                sub_table_data = []

                if parse_result.get("has_main_table"):
                    main_table_data = self.table_extractor.extract_sub_table(
                        merged_table, parse_result["main_table"]
                    )

                if parse_result.get("has_remarks"):
                    for remark in parse_result.get("remarks", []):
                        remark_content = self.table_extractor.extract_sub_table(merged_table, remark)
                        remark_data.append({
                            "region": remark,
                            "data": remark_content
                        })

                if parse_result.get("has_sub_tables"):
                    for sub_tab in parse_result.get("sub_tables", []):
                        sub_content = self.table_extractor.extract_sub_table(merged_table, sub_tab)
                        sub_table_data.append({
                            "region": sub_tab,
                            "data": sub_content
                        })

                parsed_tables.append({
                    "table_id": f"T_{idx+1:04d}_{sub_idx+1}",
                    "original_data": merged_table,
                    "parse_result": parse_result,
                    "main_table": main_table_data,
                    "remarks": remark_data,
                    "sub_tables": sub_table_data,
                    "metadata": {
                        "rows": len(merged_table),
                        "cols": max(len(r) for r in merged_table) if merged_table else 0
                    }
                })
                print(f"        子表格 {sub_idx+1}: {len(merged_table)}行, 主表: {parse_result.get('has_main_table')}, 备注: {parse_result.get('has_remarks')}, 附表: {parse_result.get('has_sub_tables')}")

        return {
            "file_name": Path(file_path).name,
            "total_pages": total_pages,
            "text_paragraphs": text_paragraphs[:50],
            "text_count": len(text_paragraphs),
            "tables": parsed_tables
        }


# ============================================================================
# Flask API
# ============================================================================

parser_system = None
tasks = {}


class ParseTask:
    def __init__(self, task_id, file_path, file_name):
        self.task_id = task_id
        self.file_path = file_path
        self.file_name = file_name
        self.status = "pending"
        self.progress = 0
        self.result = None
        self.error = None


def process_document(task):
    try:
        task.status = "processing"
        task.progress = 30

        if parser_system:
            result = parser_system.parse_document(task.file_path)
            task.result = result
            task.status = "completed"
            task.progress = 100
        else:
            task.status = "failed"
            task.error = "解析器未初始化"
    except Exception as e:
        task.status = "failed"
        task.error = str(e)
        import traceback
        traceback.print_exc()


def clean_for_json(obj):
    if obj is None:
        return None
    if isinstance(obj, (str, int, float, bool)):
        return obj
    if isinstance(obj, (list, tuple)):
        return [clean_for_json(item) for item in obj]
    if isinstance(obj, dict):
        return {str(k): clean_for_json(v) for k, v in obj.items()}
    return str(obj)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "没有上传文件"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "文件名为空"}), 400

    task_id = str(uuid.uuid4())[:8]
    saved_filename = f"{task_id}_{file.filename}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], saved_filename)
    file.save(file_path)

    task = ParseTask(task_id, file_path, file.filename)
    tasks[task_id] = task

    thread = threading.Thread(target=process_document, args=(task,))
    thread.daemon = True
    thread.start()

    return jsonify({"task_id": task_id, "status": "pending"})


@app.route('/api/task/<task_id>', methods=['GET'])
def get_task_status(task_id):
    if task_id not in tasks:
        return jsonify({"error": "任务不存在"}), 404
    task = tasks[task_id]
    return jsonify({
        "task_id": task.task_id,
        "status": task.status,
        "progress": task.progress,
        "error": task.error
    })


@app.route('/api/result/<task_id>', methods=['GET'])
def get_result(task_id):
    if task_id not in tasks:
        return jsonify({"error": "任务不存在"}), 404
    task = tasks[task_id]
    if task.status == "completed":
        return jsonify(clean_for_json(task.result))
    elif task.status == "failed":
        return jsonify({"error": task.error}), 500
    else:
        return jsonify({"status": task.status, "progress": task.progress}), 202


def init_parser():
    global parser_system
    try:
        parser_system = GeotechnicalParserSystem()
        return True
    except Exception as e:
        print(f"初始化失败: {e}")
        return False


if __name__ == '__main__':
    print("=" * 60)
    print("岩土报告智能解析系统 - 后端服务")
    print("访问地址: http://localhost:5000")
    print("=" * 60)
    with app.app_context():
        init_parser()
    app.run(debug=True, host='0.0.0.0', port=5000)