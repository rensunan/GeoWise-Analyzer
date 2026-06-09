"""
- 岩土报告智能解析系统
基于大语言模型的表格解析 - 主表/备注/附表识别
"""

import os
import re
import json
import uuid
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
from table_deduplicator import TableDeduplicator
import threading
import warnings
import hashlib

# 导入配置文件
from config import (
    FLASK_CONFIG,
    DEEPSEEK_API_KEY,
    DEEPSEEK_BASE_URL,
    DOC_PARSER_CONFIG,
    TABLE_PROCESSOR_CONFIG,
    LLM_CONFIG,
    CACHE_CONFIG,
    PATH_CONFIG
)

from table_header_alignment import get_header_aligner

warnings.filterwarnings('ignore')

from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from openai import OpenAI

app = Flask(__name__)
CORS(app)

# 使用配置文件中的设置
app.config['UPLOAD_FOLDER'] = PATH_CONFIG['uploads']
app.config['MAX_CONTENT_LENGTH'] = FLASK_CONFIG['MAX_CONTENT_LENGTH']
app.config['SECRET_KEY'] = FLASK_CONFIG['SECRET_KEY']

# 创建必要的目录
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(PATH_CONFIG['results'], exist_ok=True)

# 初始化OpenAI客户端
client = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url=DEEPSEEK_BASE_URL
)

# ============================================================================
# 文档解析器
# ============================================================================

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装")

try:
    import fitz
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("警告: PyMuPDF未安装")


class DocumentParser:
    def __init__(self):
        # 使用配置文件中的关键词
        self.content_keywords = DOC_PARSER_CONFIG['content_keywords']
        self.skip_keywords = DOC_PARSER_CONFIG['skip_keywords']
        self.min_paragraph_length = DOC_PARSER_CONFIG['min_paragraph_length']
        self.min_content_length = DOC_PARSER_CONFIG['min_content_length']
        self.skip_paragraph_length = DOC_PARSER_CONFIG['skip_paragraph_length']

    def parse_document(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        file_ext = Path(file_path).suffix.lower()
        if file_ext == '.docx':
            return self._parse_docx(file_path)
        elif file_ext == '.pdf':
            return self._parse_pdf(file_path)
        else:
            return self._parse_txt(file_path)

    def _parse_docx(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        if not DOCX_AVAILABLE:
            raise ImportError("请安装 python-docx")

        doc = Document(file_path)
        text_paragraphs = []
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if len(text) < self.min_paragraph_length:
                continue
            if any(kw in text for kw in self.skip_keywords) and len(text) < self.skip_paragraph_length:
                continue
            has_keyword = any(kw in text for kw in self.content_keywords)
            is_table_title = "表" in text and re.search(r'\d+\.\d+', text) is not None
            if has_keyword or is_table_title:
                text_paragraphs.append({
                    "id": f"P_{i+1:04d}",
                    "content": text,
                    "page": 1,
                    "is_table_title": is_table_title
                })

        return text_paragraphs, tables, 1

    def _parse_pdf(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        if not PDF_AVAILABLE:
            raise ImportError("请安装 PyMuPDF")

        doc = fitz.open(file_path)
        text_paragraphs = []
        tables = []
        para_counter = 0

        for page_num, page in enumerate(doc):
            text = page.get_text()
            for para in text.split('\n\n'):
                para = para.strip()
                if len(para) < self.min_content_length:
                    continue
                if any(kw in para for kw in self.content_keywords):
                    para_counter += 1
                    text_paragraphs.append({
                        "id": f"P_{page_num+1:04d}_{para_counter:03d}",
                        "content": para,
                        "page": page_num + 1,
                        "is_table_title": False
                    })

            page_tables = page.find_tables()
            for tab in page_tables:
                table_data = []
                for row in tab.extract():
                    row_data = [str(cell).strip() if cell else "" for cell in row]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

        doc.close()
        return text_paragraphs, tables, len(doc)

    def _parse_txt(self, file_path: str) -> Tuple[List[Dict], List[List[List[str]]], int]:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        text_paragraphs = []
        tables = []

        for i, para in enumerate(content.split('\n\n')):
            para = para.strip()
            if len(para) > self.min_content_length and any(kw in para for kw in self.content_keywords):
                text_paragraphs.append({
                    "id": f"P_{i+1:04d}",
                    "content": para,
                    "page": 1,
                    "is_table_title": False
                })

        return text_paragraphs, tables, 1


# ============================================================================
# 表格处理器
# ============================================================================

class TableProcessor:
    """表格处理器"""

    def __init__(self):
        self.header_similarity_threshold = TABLE_PROCESSOR_CONFIG['header_similarity_threshold']

    @staticmethod
    def merge_adjacent_rows(table_data: List[List[str]]) -> List[List[str]]:
        """合并完全相同的相邻行，删除多余行"""
        if not table_data:
            return table_data

        result = []
        prev_row = None

        for row in table_data:
            if prev_row is None or row != prev_row:
                result.append(row)
                prev_row = row

        return result

    @staticmethod
    def merge_adjacent_cols(table_data: List[List[str]]) -> List[List[str]]:
        """合并完全相同的相邻列，删除多余列"""
        if not table_data:
            return table_data

        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        padded = []
        for row in table_data:
            while len(row) < cols:
                row.append("")
            padded.append(row)

        cols_to_delete = [False] * cols

        for j in range(cols - 1):
            is_same = True
            for i in range(rows):
                if padded[i][j] != padded[i][j + 1]:
                    is_same = False
                    break
            if is_same:
                cols_to_delete[j + 1] = True

        new_table = []
        for i in range(rows):
            new_row = []
            for j in range(cols):
                if not cols_to_delete[j]:
                    new_row.append(padded[i][j])
            new_table.append(new_row)

        return new_table

    def split_by_header(self, table_data: List[List[str]]) -> List[List[List[str]]]:
        """按相同表头切分表格 - 相似度超过阈值即可"""
        if not table_data or len(table_data) < 2:
            return [table_data]

        header_row = table_data[0]
        header_non_empty = [(j, header_row[j]) for j in range(len(header_row)) if header_row[j]]

        if not header_non_empty:
            return [table_data]

        split_indices = [0]
        for i in range(1, len(table_data)):
            current_row = table_data[i]

            match_count = 0
            for j, header_text in header_non_empty:
                if j < len(current_row) and current_row[j] == header_text:
                    match_count += 1

            similarity = match_count / len(header_non_empty) if header_non_empty else 0

            if similarity >= self.header_similarity_threshold:
                split_indices.append(i)

        result = []
        for idx, start in enumerate(split_indices):
            end = split_indices[idx + 1] if idx + 1 < len(split_indices) else len(table_data)
            result.append(table_data[start:end])

        return result


# ============================================================================
# 大语言模型表格解析器
# ============================================================================

class LLMTableParser:
    """基于大语言模型的表格解析器"""

    def __init__(self):
        self.cache = {} if CACHE_CONFIG['enabled'] else None
        self.max_cache_size = CACHE_CONFIG['max_size']
        self.model = LLM_CONFIG['model']
        self.temperature = LLM_CONFIG['temperature']
        self.max_tokens = LLM_CONFIG['max_tokens']
        self.max_rows_for_prompt = LLM_CONFIG['max_table_rows_for_prompt']

    def _get_table_signature(self, table_data: List[List[str]]) -> str:
        """获取表格结构签名（用于判断是否为相同类型表格）"""
        if not table_data:
            return ""

        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        # 提取表头行
        header_row = table_data[0] if rows > 0 else []

        # 签名 = 行数范围 + 列数 + 表头内容哈希
        row_range = "small" if rows < 20 else "medium" if rows < 50 else "large"
        header_hash = hashlib.md5("|".join(header_row).encode()).hexdigest()[:8]

        return f"{row_range}_{cols}_{header_hash}"

    def _call_deepseek(self, table_data: List[List[str]]) -> Dict:
        """调用DeepSeek API解析表格 - 使用OpenAI SDK"""

        # 将表格转换为文本表示
        table_text = []
        for i, row in enumerate(table_data[:self.max_rows_for_prompt]):
            row_text = " | ".join([str(cell) if cell else "" for cell in row])
            table_text.append(f"行{i}: {row_text}")
        table_str = "\n".join(table_text)

        prompt = f"""你是岩土工程表格解析专家。分析表格，将每个单元格互斥地划分为"主表"、"附表"或"备注"。

        【优先级】备注 > 主表 > 附表

        【判断标准】

        1. **备注**：说明性文字
           - 包括键值对格式（如"参数: 值"）
           - 或者以"建议"、"说明"、"注"、"备注"开头的文本
           - 大段文本（>{TABLE_PROCESSOR_CONFIG['remark_max_length']}字符）或包含换行符\\n
           - 相邻列或行内容相同或类似（被切分的备注单元格或同一类型）

        2. **主表**：核心结构化数据
           - 包含岩土关键参数：孔号、孔深、土层编号、含水量、孔隙比、液限、塑限、压缩模量、黏聚力、内摩擦角等
           - 有明确的表头和数据行（数字/短文本）
           - 通常位于表格左侧/顶部，数据行数最多

        3. **附表**：仅在主表无法覆盖时才使用
           - 主表之外另起的结构化数据区域
           - 有独立的表头和数据行
           - **非必要不使用附表，优先将区域归入主表**

        【输出格式】
        {{
            "has_main_table": true/false,
            "main_table": {{
                "start_row": 0,
                "end_row": 行数-1,
                "start_col": 0,
                "end_col": 列数-1,
                "headers": [{{"col_index": 0, "start_col": 0, "end_col": 0, "header_name": "表头名"}}],
                "description": "描述"
            }},
            "has_sub_tables": true/false,
            "sub_tables": [
                {{
                    "start_row": 起始行,
                    "end_row": 结束行,
                    "start_col": 起始列,
                    "end_col": 结束列,
                    "headers": [{{"col_index": 0, "start_col": 0, "end_col": 0, "header_name": "表头名"}}],
                    "description": "描述"
                }}
            ],
            "has_remarks": true/false,
            "remarks": [
                {{
                    "start_row": 起始行,
                    "end_row": 结束行,
                    "start_col": 起始列,
                    "end_col": 结束列,
                    "content": "备注摘要"
                }}
            ]
        }}

        【规则】
        - 行号和列号从0开始
        - 每个表格必须有且只有一个主表
        - 附表是非必要的，仅在主表确实无法覆盖时才使用
        - 备注列若相邻内容相同则合并为同一备注区域

        表格数据：
        {table_str}

        只返回JSON。"""

        try:
            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "你是一个岩土工程表格解析专家，擅长识别表格结构。只返回JSON格式结果。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                stream=False
            )

            content = response.choices[0].message.content

            # 提取JSON
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                parsed = json.loads(json_match.group())
                return parsed
            else:
                return self._get_default_parse_result(table_data)

        except Exception as e:
            print(f"解析异常: {e}")
            return self._get_default_parse_result(table_data)

    def _get_default_parse_result(self, table_data: List[List[str]]) -> Dict:
        """默认解析结果（API失败时的降级方案）"""
        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        # 构建默认headers
        default_headers = []
        if rows > 0 and table_data[0]:
            for j in range(cols):
                header_name = table_data[0][j] if j < len(table_data[0]) else f"列{j}"
                default_headers.append({
                    "col_index": j,
                    "start_col": j,
                    "end_col": j,
                    "header_name": header_name
                })

        return {
            "has_main_table": True,
            "main_table": {
                "start_row": 0,
                "end_row": rows - 1,
                "start_col": 0,
                "end_col": cols - 1,
                "headers": default_headers,
                "description": "整个表格"
            },
            "has_remarks": False,
            "remarks": [],
            "has_sub_tables": False,
            "sub_tables": []
        }



    @staticmethod
    def _get_column_types(table_data, start_row, end_row, start_col, end_col):
        col_types = []
        for j in range(start_col, end_col + 1):
            numeric_count = 0
            text_count = 0
            total = 0
            for i in range(start_row + 1, end_row + 1):
                if i < len(table_data) and j < len(table_data[i]):
                    cell = str(table_data[i][j]).strip()
                    if cell:
                        total += 1
                        try:
                            float(cell.replace(",", "").replace(chr(0xff0c), ""))
                            numeric_count += 1
                        except ValueError:
                            text_count += 1
            if total == 0:
                col_types.append("empty")
            elif numeric_count / total >= 0.8:
                col_types.append("numeric")
            elif text_count / total >= 0.8:
                col_types.append("text")
            else:
                col_types.append("mixed")
        return col_types

    @staticmethod
    def _row_matches_types(row, col_types, start_col):
        for j, col_type in enumerate(col_types):
            if col_type in ("mixed", "empty"):
                continue
            actual_col = start_col + j
            if actual_col >= len(row):
                continue
            cell = str(row[actual_col]).strip()
            if not cell:
                continue
            if col_type == "numeric":
                try:
                    float(cell.replace(",", "").replace(chr(0xff0c), ""))
                except ValueError:
                    return False
            elif col_type == "text":
                try:
                    float(cell.replace(",", "").replace(chr(0xff0c), ""))
                    return False
                except ValueError:
                    pass
        return True

    def _trim_by_consistency(self, table_data, adapted):
        if not adapted.get("has_main_table") or not adapted.get("main_table"):
            return adapted
        mt = adapted["main_table"]
        if mt["end_row"] <= mt["start_row"]:
            return adapted
        col_types = self._get_column_types(
            table_data, mt["start_row"], mt["end_row"],
            mt["start_col"], mt["end_col"]
        )
        if all(ct in ("mixed", "empty") for ct in col_types):
            return adapted
        for i in range(mt["start_row"] + 1, mt["end_row"] + 1):
            row = table_data[i]
            if not self._row_matches_types(row, col_types, mt["start_col"]):
                mt["end_row"] = i - 1
                print("  [consistency] main stopped at row %d, type mismatch at row %d"
                      % (mt["end_row"], i))
                break
        return adapted





    def _realign_headers(self, cached_result, table_data):
        import copy, re as _re
        if not table_data:
            return None
        total_rows = len(table_data)
        adapted = copy.deepcopy(cached_result)

        def _norm(cell):
            if not cell:
                return ''
            return _re.sub(r'\s+', '', str(cell))

        def _sim(raw, row):
            nh = [_norm(c) for c in raw if _norm(c)]
            nr = [_norm(c) for c in row if _norm(c)]
            if not nh or not nr:
                return 0.0
            m = 0
            for hc in nh:
                for rc in nr:
                    if hc in rc or rc in hc:
                        m += 1
                        break
            return m / len(nh)

        # 1. main header
        raw_h = cached_result.get('_raw_header', [])
        if not raw_h and adapted.get('main_table'):
            raw_h = [h.get('header_name','') for h in adapted['main_table'].get('headers',[])]
        if not raw_h:
            print('  [header realign] no header data, giving up')
            return None

        best_r, best_s = -1, 0.0
        for r in range(total_rows):
            s = _sim(raw_h, table_data[r])
            if s > best_s:
                best_s, best_r = s, r

        if best_s < 0.5:
            print('  [header realign] best row %d score %.2f < 0.5, giving up' % (best_r, best_s))
            print('  [header realign] cached: %s' % [_norm(c) for c in raw_h[:8]])
            print('  [header realign] row 0:  %s' % [_norm(c) for c in table_data[0][:10]])
            return None

        mt = adapted.get('main_table')
        if mt and best_r != mt.get('start_row', 0):
            shift = best_r - mt['start_row']
            print('  [header realign] main row %d -> %d (score %.2f)' % (mt['start_row'], best_r, best_s))
            mt['start_row'] = best_r
            mt['end_row'] = min(total_rows - 1, mt['end_row'] + shift)

        # 2. sub headers
        if adapted.get('has_sub_tables'):
            for st in adapted.get('sub_tables', []):
                sraw = st.get('_raw_header', [])
                if not sraw:
                    sraw = [h.get('header_name','') for h in st.get('headers',[])]
                if not sraw:
                    continue
                br, bs = -1, 0.0
                for r in range(total_rows):
                    if r == best_r:
                        continue
                    s = _sim(sraw, table_data[r])
                    if s > bs:
                        bs, br = s, r
                if bs < 0.5:
                    print('  [header realign] sub not found (best %.2f), giving up' % bs)
                    return None
                if br != st.get('start_row', 0):
                    print('  [header realign] sub row %d -> %d (score %.2f)' % (st['start_row'], br, bs))
                    shift = br - st['start_row']
                    st['start_row'] = br
                    st['end_row'] = min(total_rows - 1, st['end_row'] + shift)

        return adapted


    def _adapt_cached_result(self, cached_result, table_data):
        """自适应修正缓存结果中的行列索引，使其适配当前表格的实际尺寸。"""
        import copy
        adapted = copy.deepcopy(cached_result)
        total_rows = len(table_data)
        total_cols = max(len(r) for r in table_data) if table_data else 0
        changes = []

        def _clamp(value, lo, hi):
            return max(lo, min(value, hi))

        # 1. 主表：clamp + 自适应扩展
        if adapted.get("has_main_table") and adapted.get("main_table"):
            mt = adapted["main_table"]
            old_sr, old_er = mt["start_row"], mt["end_row"]
            old_sc, old_ec = mt["start_col"], mt["end_col"]

            mt["start_row"] = _clamp(mt["start_row"], 0, total_rows - 1)
            mt["start_col"] = _clamp(mt["start_col"], 0, total_cols - 1)
            mt["end_col"]   = _clamp(mt["end_col"], mt["start_col"], total_cols - 1)

            if adapted.get("has_sub_tables") and adapted["sub_tables"]:
                first_sub_start = min(st["start_row"] for st in adapted["sub_tables"])
                mt["end_row"] = _clamp(max(mt["end_row"], first_sub_start - 1), mt["start_row"], total_rows - 1)
            else:
                mt["end_row"] = total_rows - 1

            if old_sr != mt["start_row"] or old_er != mt["end_row"]:
                changes.append("main row {}->{} -> {}->{}".format(old_sr, old_er, mt['start_row'], mt['end_row']))
            if old_sc != mt["start_col"] or old_ec != mt["end_col"]:
                changes.append("main col {}->{} -> {}->{}".format(old_sc, old_ec, mt['start_col'], mt['end_col']))

        # 2. 附表：clamp + 排除与主表重叠的
        if adapted.get("has_sub_tables"):
            valid = []
            mt = adapted.get("main_table") if adapted.get("has_main_table") else None
            for st in adapted["sub_tables"]:
                st["start_row"] = _clamp(st["start_row"], 0, total_rows - 1)
                st["end_row"]   = _clamp(st["end_row"], st["start_row"], total_rows - 1)
                st["start_col"] = _clamp(st["start_col"], 0, total_cols - 1)
                st["end_col"]   = _clamp(st["end_col"], st["start_col"], total_cols - 1)

                if mt:
                    if st["start_row"] >= mt["start_row"] and st["end_row"] <= mt["end_row"]:
                        changes.append("sub-table row {}-{} overlaps main, removed".format(st['start_row'], st['end_row']))
                        continue
                if st["end_row"] >= st["start_row"] and st["end_col"] >= st["start_col"]:
                    valid.append(st)
            adapted["sub_tables"] = valid
            adapted["has_sub_tables"] = len(valid) > 0

        # 3. 备注：clamp + 排除与主表重叠的
        if adapted.get("has_remarks"):
            valid = []
            mt = adapted.get("main_table") if adapted.get("has_main_table") else None
            for r in adapted["remarks"]:
                r["start_row"] = _clamp(r["start_row"], 0, total_rows - 1)
                r["end_row"]   = _clamp(r["end_row"], r["start_row"], total_rows - 1)
                r["start_col"] = _clamp(r["start_col"], 0, total_cols - 1)
                r["end_col"]   = _clamp(r["end_col"], r["start_col"], total_cols - 1)

                if mt:
                    if r["start_row"] >= mt["start_row"] and r["end_row"] <= mt["end_row"]:
                        changes.append("remark row {}-{} overlaps main, removed".format(r['start_row'], r['end_row']))
                        continue
                if r["end_row"] >= r["start_row"] and r["end_col"] >= r["start_col"]:
                    valid.append(r)
            adapted["remarks"] = valid
            adapted["has_remarks"] = len(valid) > 0

        # 4. Headers 中的 col_index 也做 clamp
        for section in ["main_table", "sub_tables"]:
            items = []
            if section == "main_table" and adapted.get("main_table"):
                items = [adapted["main_table"]]
            elif section == "sub_tables":
                items = adapted.get("sub_tables", [])
            for item in items:
                for h in item.get("headers", []):
                    h["col_index"] = _clamp(h.get("col_index", 0), 0, total_cols - 1)
                    h["start_col"] = _clamp(h.get("start_col", 0), 0, total_cols - 1)
                    h["end_col"]   = _clamp(h.get("end_col", 0), 0, total_cols - 1)

        if changes:
            print("  [cache adapt] table {}x{}, {} change(s):".format(total_rows, total_cols, len(changes)))
            for c in changes:
                print("    - {}".format(c))
        else:
            print("  [cache adapt] table {}x{}, no changes needed".format(total_rows, total_cols))

        adapted = self._trim_by_consistency(table_data, adapted)
        return adapted

    def parse(self, table_data: List[List[str]]) -> Dict:
        """解析表格，返回主表/备注/附表的位置信息"""
        if not table_data:
            return {
                "has_main_table": False,
                "main_table": None,
                "has_remarks": False,
                "remarks": [],
                "has_sub_tables": False,
                "sub_tables": []
            }

        # 获取表格签名
        signature = self._get_table_signature(table_data)

        # 检查缓存
        if self.cache is not None and signature in self.cache:
            cached = self.cache[signature]
            realigned = self._realign_headers(cached, table_data)
            if realigned is not None:
                print(f"  [缓存命中] 使用相同结构表格的解析结果")
                return self._adapt_cached_result(realigned, table_data)
            else:
                print(f"  [缓存命中但表头检索失败] 重新调用大模型")

        # 调用大模型
        print(f"  [调用DeepSeek] 解析表格 {len(table_data)}行 x {len(table_data[0]) if table_data else 0}列")
        result = self._call_deepseek(table_data)

        # 缓存结果
        if self.cache is not None:
            if len(self.cache) >= self.max_cache_size:
                oldest_key = next(iter(self.cache))
                del self.cache[oldest_key]
            self.cache[signature] = result
            if table_data and len(table_data) > 0:
                result["_raw_header"] = table_data[0]

        return result


# ============================================================================
# 表格提取器
# ============================================================================

class TableExtractor:
    """根据LLM解析结果提取子表格"""

    def __init__(self):
        self.remark_keywords = TABLE_PROCESSOR_CONFIG['remark_keywords']
        self.remark_max_length = TABLE_PROCESSOR_CONFIG['remark_max_length']
        self.remark_similarity_threshold = TABLE_PROCESSOR_CONFIG['remark_similarity_threshold']
        self.remark_min_length_for_merge = TABLE_PROCESSOR_CONFIG['remark_min_length_for_merge']

    @staticmethod
    def extract_sub_table(table_data: List[List[str]], region: Dict) -> List[List[str]]:
        """提取指定区域的子表格"""
        if not region:
            return []

        start_row = region.get("start_row", 0)
        end_row = region.get("end_row", len(table_data) - 1)
        start_col = region.get("start_col", 0)
        end_col = region.get("end_col", 0)

        start_row = max(0, min(start_row, len(table_data) - 1))
        end_row = max(start_row, min(end_row, len(table_data) - 1))
        start_col = max(0, start_col)
        end_col = max(start_col, end_col)

        result = []
        for i in range(start_row, end_row + 1):
            if i < len(table_data):
                row = table_data[i]
                if start_col < len(row):
                    result.append(row[start_col:min(end_col + 1, len(row))])
                else:
                    result.append([])

        return result

    def fix_remark_boundary(self, table_data: List[List[str]], parse_result: Dict) -> Dict:
        """修正备注边界：将包含备注内容的列从附表中移除"""
        remarks = parse_result.get("remarks", [])
        if not remarks:
            return parse_result

        sub_tables = parse_result.get("sub_tables", [])
        if not sub_tables:
            return parse_result

        for sub_tab in sub_tables:
            sub_start_col = sub_tab.get("start_col", 0)
            sub_start_row = sub_tab.get("start_row", 0)

            if sub_start_row < len(table_data) and sub_start_col < len(table_data[sub_start_row]):
                cell_content = table_data[sub_start_row][sub_start_col]

                is_remark = False
                if cell_content:
                    if '\n' in cell_content:
                        is_remark = True
                    elif len(cell_content) > self.remark_max_length:
                        is_remark = True
                    elif any(keyword in cell_content for keyword in self.remark_keywords):
                        is_remark = True

                if is_remark:
                    new_start_col = sub_start_col + 1
                    while new_start_col < len(table_data[sub_start_row]):
                        next_cell = table_data[sub_start_row][new_start_col]
                        is_next_remark = False
                        if next_cell:
                            if '\n' in next_cell:
                                is_next_remark = True
                            elif len(next_cell) > self.remark_max_length:
                                is_next_remark = True
                            elif any(keyword in next_cell for keyword in self.remark_keywords):
                                is_next_remark = True

                        if not is_next_remark:
                            break
                        new_start_col += 1

                    sub_tab["start_col"] = new_start_col

                    old_headers = sub_tab.get("headers", [])
                    new_headers = []
                    for header in old_headers:
                        if header.get("col_index", 0) >= new_start_col:
                            header["col_index"] = header["col_index"] - (sub_start_col - new_start_col)
                            new_headers.append(header)
                    sub_tab["headers"] = new_headers

                    remark_found = False
                    for remark in remarks:
                        if remark.get("start_row", 0) <= sub_start_row <= remark.get("end_row", 0):
                            remark["start_col"] = min(remark.get("start_col", sub_start_col), sub_start_col)
                            remark["end_col"] = max(remark.get("end_col", sub_start_col), sub_start_col)
                            remark_found = True
                            break

                    if not remark_found:
                        remarks.append({
                            "start_row": sub_start_row,
                            "end_row": sub_start_row,
                            "start_col": sub_start_col,
                            "end_col": sub_start_col,
                            "content": "备注内容"
                        })

        parse_result["remarks"] = remarks
        parse_result["sub_tables"] = sub_tables

        return parse_result

    def _deduplicate_remark_content(self, remark_content: List[List[str]]) -> List[List[str]]:
        """对单个备注块的内容进行去重"""
        if not remark_content:
            return remark_content

        deduplicated_rows = []
        for row in remark_content:
            unique_cells = []
            prev_cell = None
            for cell in row:
                cell_clean = cell.strip() if cell else ""
                if cell_clean and cell_clean != prev_cell:
                    unique_cells.append(cell)
                    prev_cell = cell_clean
                elif not cell_clean:
                    unique_cells.append(cell)
            deduplicated_rows.append(unique_cells)

        unique_rows = []
        seen_row_strs = set()
        for row in deduplicated_rows:
            row_str = " | ".join([str(cell) for cell in row if cell and cell.strip()])
            if row_str and row_str not in seen_row_strs:
                seen_row_strs.add(row_str)
                unique_rows.append(row)
            elif not row_str and not any(cell for cell in row):
                continue

        if len(unique_rows) > 1:
            merged_rows = []
            current_merged = []
            for row in unique_rows:
                row_text = " ".join([str(cell) for cell in row if cell and cell.strip()])
                if current_merged:
                    last_text = " ".join([str(cell) for cell in current_merged[-1] if cell and cell.strip()])
                    if row_text and not re.match(r'^\d+[、.]', row_text):
                        current_merged[-1] = [current_merged[-1][0] + " " + row_text] if current_merged[-1] else [row_text]
                    else:
                        current_merged.append(row)
                else:
                    current_merged.append(row)
            unique_rows = current_merged

        return unique_rows

    def _deduplicate_remarks_list(self, remarks: List[Dict]) -> List[Dict]:
        """对备注列表进行全局去重"""
        if not remarks:
            return remarks

        remark_texts = []
        for remark in remarks:
            data = remark.get("data", [])
            text_parts = []
            seen_phrases = set()
            for row in data:
                for cell in row:
                    if cell and cell.strip():
                        cell_clean = cell.strip()
                        if cell_clean not in seen_phrases:
                            seen_phrases.add(cell_clean)
                            text_parts.append(cell_clean)
            full_text = " ".join(text_parts)
            remark_texts.append(full_text)

        unique_remarks = []
        unique_texts = []

        for i, remark in enumerate(remarks):
            text = remark_texts[i]
            if not text:
                continue

            is_duplicate = False
            for j, existing_text in enumerate(unique_texts):
                if text in existing_text or existing_text in text:
                    is_duplicate = True
                    if len(text) > len(existing_text):
                        unique_texts[j] = text
                        unique_remarks[j] = remark
                    break

                min_len = min(len(text), len(existing_text))
                if min_len > self.remark_min_length_for_merge:
                    common = self._longest_common_substring(text, existing_text)
                    if common / min_len > self.remark_similarity_threshold:
                        is_duplicate = True
                        break

            if not is_duplicate:
                unique_texts.append(text)
                unique_remarks.append(remark)

        return unique_remarks

    def _longest_common_substring(self, s1: str, s2: str) -> int:
        """计算两个字符串的最长公共子串长度"""
        if not s1 or not s2:
            return 0

        len1, len2 = len(s1), len(s2)
        if len1 > len2:
            s1, s2 = s2, s1
            len1, len2 = len2, len1

        max_len = 0
        for i in range(len1):
            for j in range(len2):
                k = 0
                while i + k < len1 and j + k < len2 and s1[i + k] == s2[j + k]:
                    k += 1
                if k > max_len:
                    max_len = k
        return max_len


# ============================================================================
# 二次大模型解析器
# ============================================================================

class SecondaryLLMParser:
    """二次大模型解析器 - 处理合并表头（重复行/列已通过硬编码去除）"""

    def __init__(self):
        self.model = LLM_CONFIG['model']
        self.temperature = LLM_CONFIG['temperature']
        self.max_tokens = LLM_CONFIG['secondary_max_tokens']
        self.max_rows_for_prompt = LLM_CONFIG['max_table_rows_for_prompt']

    def _call_deepseek_for_optimization(self, table_data: List[List[str]], table_name: str) -> Dict:
        """调用DeepSeek API优化表格 - 只处理表头合并和单位对齐"""
        table_text = []
        for i, row in enumerate(table_data[:self.max_rows_for_prompt]):
            row_text = " | ".join([str(cell) if cell else "" for cell in row])
            table_text.append(f"行{i}: {row_text}")
        table_str = "\n".join(table_text)

        # 修改提示词，去除重复行/列的处理要求
        prompt = f"""你是一个岩土工程表格优化专家。请分析以下表格（{table_name}），并进行优化处理。

        【优化规则】

        1. **拆分复合表头列**
           - 识别表头单元格中包含多个独立概念的情况（如"A B"、"A-B"、"A\nB"、"A和B"等）
           - 常见复合表头模式：
             * "钻孔编号孔深" → ["钻孔编号", "孔深"]
             * "杆塔号 塔型" → ["杆塔号", "塔型"]
             * "地层编号-时代成因" → ["地层编号", "时代成因"]
           - 拆分后，对应数据列按相同分隔符拆分成多个值

        2. **跨行数据填充（合并分散数据）**
           - 拆分复合列后，检查新列的数据分布情况
           - 如果某一列满足以下条件：
             * 大部分行（>70%）为空值
             * 少数行（<30%）有非空值
             * 这些有值的行，其他列的数据与某正常行重复
           - 则执行跨行填充：
             * 提取这些有值行的内容作为"公共值"
             * 将公共值填充到该列的所有空行
             * 删除这些有值的特殊行（因为数据已被提取）

           示例：
           拆分前：
             ["孔号孔深", "含水率"]
             ["ZK1", "25.3"]
             ["2.5m", "25.3"]    # 特殊行：只有孔深，孔号为空
             ["ZK2", "26.1"]

           拆分后：
             ["孔号", "孔深", "含水率"]
             ["ZK1", "", "25.3"]
             ["", "2.5m", "25.3"]    # 特殊行
             ["ZK2", "", "26.1"]

           填充后：
             ["孔号", "孔深", "含水率"]
             ["ZK1", "2.5m", "25.3"]
             ["ZK2", "2.5m", "26.1"]
             (删除特殊行)

        3. **表头标准化**
           - 统一表头命名格式
           - 如果数据中包含单位（如"25.3%"、"14kN/m³"），将单位提取到表头：`"数据值(单位)"`
           - 示例：
             * 列数据["14", "16", "15"]，表头"γ" → 改为"γ(kN/m³)"
             * 列数据["25.3%", "26.1%", "24.8%"]，表头"含水率" → 改为"含水率(%)"

        【执行流程】
        - 步骤1：如果表头包含复合概念，先拆分列
        - 步骤2：拆分后，检查是否需要跨行填充
        - 步骤3：最后进行表头标准化（添加单位）

        【输出格式】
        {{
            "has_optimization": true/false,
            "optimization_type": ["split_headers", "cross_row_fill", "standardize_units"],
            "optimized_table": [
                ["优化后的表头1", "列2", "列3"],
                ["数据1", "值2", "值3"]
            ],
            "explanation": "优化说明"
        }}

        【原始表格数据】
        {table_str}

        重要：
        - 如果表格无需优化，optimized_table返回原始表格
        - has_optimization为true时，optimized_table必须是优化后的完整表格
        - 只返回JSON，不要有其他内容"""

        try:
            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system",
                     "content": "你是一个岩土工程表格优化专家。只返回JSON格式结果，必须包含optimized_table字段。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                stream=False
            )

            content = response.choices[0].message.content

            # 提取JSON
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                parsed = json.loads(json_match.group())
                return parsed
            else:
                return self._get_default_result(table_data)

        except Exception as e:
            print(f"二次解析异常: {e}")
            return self._get_default_result(table_data)

    def _get_default_result(self, table_data: List[List[str]]) -> Dict:
        return {
            "has_optimization": False,
            "optimization_type": [],
            "optimized_table": table_data,
            "explanation": "无需优化"
        }

    def optimize_table(self, table_data: List[List[str]], table_name: str) -> Dict:
        """优化表格 - 先硬编码去重，再调用LLM处理表头合并"""
        if not table_data or len(table_data) < 1:
            return {
                "has_optimization": False,
                "optimized_table": table_data,
                "explanation": "表格为空"
            }

        print(f"\n{'=' * 80}")
        print(f"[二次解析] 开始处理: {table_name}")
        print(f"{'=' * 80}")

        # 显示原始表格信息
        print(f"\n[原始表格信息]")
        print(f"  行数: {len(table_data)}")
        print(f"  列数: {max(len(row) for row in table_data) if table_data else 0}")

        # 显示原始表格内容（前10行）
        print(f"\n[原始表格内容] (显示前10行)")
        print(f"-" * 80)
        for i, row in enumerate(table_data[:10]):
            row_str = " | ".join([str(cell)[:30] + "..." if len(str(cell)) > 30 else str(cell) for cell in row])
            print(f"  行{i:3d}: {row_str}")
        if len(table_data) > 10:
            print(f"  ... (还有 {len(table_data) - 10} 行)")
        print(f"-" * 80)

        print(f"\n[硬编码去重处理]")
        # 使用硬编码去重（保护表头）
        dedup_result = TableDeduplicator.deduplicate_with_header_protection(table_data)
        deduplicated_table = dedup_result["optimized_table"]

        if dedup_result["has_optimization"]:
            print(f"  ✓ 发现重复数据")
            print(f"    删除重复行: {dedup_result['stats']['rows_removed']} 行")
            print(f"    删除重复列: {dedup_result['stats']['cols_removed']} 列")
            print(
                f"    表格大小变化: {dedup_result['stats']['original_rows']}x{dedup_result['stats']['original_cols']} → {dedup_result['stats']['final_rows']}x{dedup_result['stats']['final_cols']}")
        else:
            print(f"  ✓ 未发现重复行或重复列")

        # 显示去重后的表格内容
        print(f"\n[去重后表格内容] (输入LLM前的表格)")
        print(f"-" * 80)
        for i, row in enumerate(deduplicated_table[:10]):
            row_str = " | ".join([str(cell)[:30] + "..." if len(str(cell)) > 30 else str(cell) for cell in row])
            print(f"  行{i:3d}: {row_str}")
        if len(deduplicated_table) > 10:
            print(f"  ... (还有 {len(deduplicated_table) - 10} 行)")
        print(f"-" * 80)

        # 显示将要发送给LLM的数据统计
        print(f"\n[LLM输入数据统计]")
        print(f"  行数: {len(deduplicated_table)}")
        print(f"  列数: {max(len(row) for row in deduplicated_table) if deduplicated_table else 0}")
        print(f"  总单元格数: {sum(len(row) for row in deduplicated_table)}")

        print(f"\n[调用DeepSeek API] 优化表头合并...")
        result = self._call_deepseek_for_optimization(deduplicated_table, table_name)

        # 显示LLM优化结果
        if result.get("has_optimization"):
            print(f"  ✓ LLM优化完成")
            print(f"    优化类型: {result.get('optimization_type', [])}")
            print(f"    优化说明: {result.get('explanation', '')}")

            # 显示优化后的表格预览
            optimized_table = result.get("optimized_table", [])
            print(f"\n[LLM优化后表格预览] (前5行)")
            print(f"-" * 80)
            for i, row in enumerate(optimized_table[:5]):
                row_str = " | ".join([str(cell)[:30] + "..." if len(str(cell)) > 30 else str(cell) for cell in row])
                print(f"  行{i:3d}: {row_str}")
            if len(optimized_table) > 5:
                print(f"  ... (还有 {len(optimized_table) - 5} 行)")
            print(f"-" * 80)
        else:
            print(f"  ℹ LLM未发现需要优化的内容")

        # 合并去重信息和LLM优化信息
        if dedup_result["has_optimization"]:
            result["deduplication_stats"] = dedup_result["stats"]
            if not result.get("has_optimization", False):
                result["has_optimization"] = True
                result["optimization_type"].append("duplicate_rows_cols")

        print(f"\n[二次解析完成] {table_name}")
        print(f"{'=' * 80}\n")

        return result


# ============================================================================
# 主解析系统
# ============================================================================

class GeotechnicalParserSystem:
    def __init__(self):
        print("=" * 60)
        print("岩土报告智能解析系统")
        print("基于大语言模型的表格解析 - 主表/备注/附表识别")
        print("=" * 60)

        print("\n[1] 初始化文档解析器...")
        self.doc_parser = DocumentParser()

        print("\n[2] 初始化表格处理器...")
        self.table_processor = TableProcessor()

        print("\n[3] 初始化大语言模型解析器...")
        self.llm_parser = LLMTableParser()

        print("\n[4] 初始化二次大语言模型解析器...")
        self.secondary_parser = SecondaryLLMParser()

        print("\n[5] 初始化表格提取器...")
        self.table_extractor = TableExtractor()

        print("\n[6] 初始化表头对齐器...")
        self.header_aligner = get_header_aligner()

        print("\n✓ 系统初始化完成！")

    def parse_document(self, file_path: str) -> Dict:
        print(f"\n[解析文档] {file_path}")

        print("  步骤1: 提取文本段落和表格...")
        text_paragraphs, raw_tables, total_pages = self.doc_parser.parse_document(file_path)
        print(f"    文本段落: {len(text_paragraphs)}")
        print(f"    原始表格: {len(raw_tables)}")

        print("  步骤2: 按表头切分表格...")
        print("  步骤3: 对每个子表格合并相邻相同行列...")
        print("  步骤4: 大语言模型解析表格结构...")

        parsed_tables = []
        for idx, table_data in enumerate(raw_tables):
            print(f"    处理表格 T_{idx + 1:04d}: 原始 {len(table_data)}行")

            split_tables = self.table_processor.split_by_header(table_data)
            print(f"      切分为 {len(split_tables)} 个独立表格")

            for sub_idx, sub_table in enumerate(split_tables):
                merged_rows = self.table_processor.merge_adjacent_rows(sub_table)
                merged_table = self.table_processor.merge_adjacent_cols(merged_rows)

                cleaned_table, groundwater_remarks = self._extract_groundwater_rows(merged_table)

                parse_result = self.llm_parser.parse(cleaned_table)
                parse_result = self.table_extractor.fix_remark_boundary(cleaned_table, parse_result)

                main_table_data = []
                remark_data = []
                sub_table_data = []

                if parse_result.get("has_main_table"):
                    main_table_data = self.table_extractor.extract_sub_table(
                        cleaned_table, parse_result["main_table"]
                    )
                    llm_headers = parse_result.get("main_table", {}).get("headers", [])
                    header_names = [h.get("header_name", "") for h in llm_headers]
                    main_table_data = self._fill_consistent_columns(main_table_data, header_names)

                if groundwater_remarks:
                    remark_data.extend(groundwater_remarks)

                if parse_result.get("has_remarks"):
                    for remark in parse_result.get("remarks", []):
                        remark_content = self.table_extractor.extract_sub_table(cleaned_table, remark)
                        deduplicated_content = self.table_extractor._deduplicate_remark_content(remark_content)
                        remark_data.append({
                            "region": remark,
                            "data": deduplicated_content
                        })

                remark_data = self.table_extractor._deduplicate_remarks_list(remark_data)

                if parse_result.get("has_sub_tables"):
                    for sub_tab in parse_result.get("sub_tables", []):
                        sub_content = self.table_extractor.extract_sub_table(cleaned_table, sub_tab)
                        st_headers = sub_tab.get("headers", [])
                        st_header_names = [h.get("header_name", "") for h in st_headers]
                        sub_content = self._fill_consistent_columns(sub_content, st_header_names)
                        sub_table_data.append({
                            "region": sub_tab,
                            "data": sub_content
                        })

                parsed_tables.append({
                    "table_id": f"T_{idx + 1:04d}_{sub_idx + 1}",
                    "original_data": cleaned_table,
                    "parse_result": parse_result,
                    "main_table": main_table_data,
                    "remarks": remark_data,
                    "sub_tables": sub_table_data,
                    "secondary_optimized": None,
                    "metadata": {
                        "rows": len(merged_table),
                        "cols": max(len(r) for r in merged_table) if merged_table else 0
                    }
                })
                print(
                    f"        子表格 {sub_idx + 1}: {len(merged_table)}行, 主表: {parse_result.get('has_main_table')}, 备注: {parse_result.get('has_remarks')}, 附表: {parse_result.get('has_sub_tables')}")

        return {
            "file_name": Path(file_path).name,
            "total_pages": total_pages,
            "text_paragraphs": text_paragraphs[:50],
            "text_count": len(text_paragraphs),
            "tables": parsed_tables
        }


    def _extract_groundwater_rows(self, table_data):
        import copy
        groundwater_keywords = TABLE_PROCESSOR_CONFIG["groundwater_keywords"]
        max_zone_cells = 4

        if not table_data or len(table_data) < 2:
            return table_data, []

        cleaned = copy.deepcopy(table_data)

        header_row = cleaned[0] if cleaned else []
        for cell in header_row:
            if self._matches_groundwater(cell, groundwater_keywords):
                print("  [groundwater] header has keyword: " + cell + ", skip")
                return cleaned, []

        groundwater_remarks = []

        for row_idx in range(1, len(cleaned)):
            row = cleaned[row_idx]
            col_cursor = 0
            while col_cursor < len(row):
                cell = row[col_cursor].strip() if row[col_cursor] else ""
                if not cell:
                    col_cursor += 1
                    continue

                if self._matches_groundwater(cell, groundwater_keywords):
                    label_text = cell
                    zone_start = col_cursor
                    zone_end = col_cursor
                    value_parts = []
                    cells_in_zone = 1

                    next_col = col_cursor + 1
                    while next_col < len(row) and cells_in_zone < max_zone_cells:
                        next_cell = row[next_col].strip() if row[next_col] else ""
                        if next_cell == label_text:
                            zone_end = next_col
                            cells_in_zone += 1
                            next_col += 1
                        else:
                            break

                    if next_col < len(row) and cells_in_zone < max_zone_cells:
                        next_cell = row[next_col].strip() if row[next_col] else ""
                        if next_cell and not self._matches_groundwater(next_cell, groundwater_keywords):
                            value_parts.append(next_cell)
                            zone_end = next_col

                    if value_parts:
                        remark_text = label_text + "：" + " ".join(value_parts)
                    else:
                        remark_text = label_text
                    groundwater_remarks.append({
                        "region": {
                            "start_row": row_idx, "end_row": row_idx,
                            "start_col": zone_start, "end_col": zone_end,
                            "content": remark_text
                        },
                        "data": [[remark_text]]
                    })

                    print("  [groundwater] extracted row" + str(row_idx) + " col" + str(zone_start) + "-" + str(zone_end) + ": " + remark_text)

                    for cc in range(zone_start, zone_end + 1):
                        if cc < len(cleaned[row_idx]):
                            cleaned[row_idx][cc] = ""

                    col_cursor = zone_end + 1
                else:
                    col_cursor += 1

        if groundwater_remarks:
            print("  [groundwater] total " + str(len(groundwater_remarks)) + " remarks extracted")
        else:
            print("  [groundwater] none found")

        return cleaned, groundwater_remarks

    @staticmethod
    def _matches_groundwater(cell_text, keywords):
        if not cell_text:
            return False
        cell_clean = cell_text.strip()
        cell_clean = cell_clean.replace(chr(10), "").replace(chr(13), "")
        for kw in keywords:
            if cell_clean.startswith(kw):
                return True
        return False

    @staticmethod
    def _fill_consistent_columns(table_data, header_names=None):
        if not table_data or len(table_data) < 2:
            return table_data

        rows = len(table_data)
        cols = max(len(row) for row in table_data)

        for row in table_data:
            while len(row) < cols:
                row.append("")

        data_start = 1
        if header_names and len(header_names) > 0:
            header_names_set = set(h.strip() for h in header_names if h and h.strip())
            last_header_row = -1
            for r in range(rows):
                row_cells = [c.strip() for c in table_data[r] if c and c.strip()]
                if not row_cells:
                    continue
                matches = sum(1 for c in row_cells if c in header_names_set)
                if matches > 0:
                    last_header_row = r
            if last_header_row >= 0:
                data_start = last_header_row + 1

        if data_start >= rows:
            return table_data

        print("  [fill] data starts at row " + str(data_start) + " (total " + str(rows) + " rows)")

        for col in range(cols):
            values = set()
            for row_idx in range(data_start, rows):
                cell = table_data[row_idx][col].strip() if col < len(table_data[row_idx]) else ""
                if cell:
                    values.add(cell)

            if len(values) == 1:
                fill_value = values.pop()
                filled_count = 0
                for row_idx in range(data_start, rows):
                    if col < len(table_data[row_idx]):
                        cell = table_data[row_idx][col].strip()
                        if not cell:
                            table_data[row_idx][col] = fill_value
                            filled_count += 1
                if filled_count > 0:
                    print("  [fill] col " + str(col) + ": filled " + str(filled_count) + " cells with '" + fill_value + "'")

        return table_data


# ============================================================================
# Flask API
# ============================================================================

parser_system = None
tasks = {}


class ParseTask:
    def __init__(self, task_id, file_path, file_name):
        self.task_id = task_id
        self.file_path = file_path
        self.file_name = file_name
        self.status = "pending"
        self.progress = 0
        self.result = None
        self.error = None


def process_document(task):
    try:
        task.status = "processing"
        task.progress = 30

        if parser_system:
            result = parser_system.parse_document(task.file_path)
            task.result = result
            task.status = "completed"
            task.progress = 100

            try:
                os.remove(task.file_path)
                print(f"已删除临时文件: {task.file_path}")
            except Exception as e:
                print(f"删除文件失败: {e}")
    except Exception as e:
        task.status = "failed"
        task.error = str(e)


def clean_for_json(obj):
    if obj is None:
        return None
    if isinstance(obj, (str, int, float, bool)):
        return obj
    if isinstance(obj, (list, tuple)):
        return [clean_for_json(item) for item in obj]
    if isinstance(obj, dict):
        return {str(k): clean_for_json(v) for k, v in obj.items()}
    return str(obj)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "没有上传文件"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "文件名为空"}), 400

    task_id = str(uuid.uuid4())[:8]
    saved_filename = f"{task_id}_{file.filename}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], saved_filename)
    file.save(file_path)

    task = ParseTask(task_id, file_path, file.filename)
    tasks[task_id] = task

    thread = threading.Thread(target=process_document, args=(task,))
    thread.daemon = True
    thread.start()

    return jsonify({"task_id": task_id, "status": "pending"})


@app.route('/api/task/<task_id>', methods=['GET'])
def get_task_status(task_id):
    if task_id not in tasks:
        return jsonify({"error": "任务不存在"}), 404
    task = tasks[task_id]
    return jsonify({
        "task_id": task.task_id,
        "status": task.status,
        "progress": task.progress,
        "error": task.error
    })


@app.route('/api/result/<task_id>', methods=['GET'])
def get_result(task_id):
    if task_id not in tasks:
        return jsonify({"error": "任务不存在"}), 404
    task = tasks[task_id]
    if task.status == "completed":
        return jsonify(clean_for_json(task.result))
    elif task.status == "failed":
        return jsonify({"error": task.error}), 500
    else:
        return jsonify({"status": task.status, "progress": task.progress}), 202


@app.route('/api/optimize/<task_id>/<int:table_index>', methods=['POST'])
def optimize_single_table(task_id, table_index):
    """对单个表格进行二次优化"""
    if task_id not in tasks:
        return jsonify({"error": "任务不存在"}), 404

    task = tasks[task_id]
    if task.status != "completed" or not task.result:
        return jsonify({"error": "任务未完成"}), 400

    try:
        data = request.get_json()
        table_name = data.get('table_name', f'表格{table_index}')

        tables = task.result.get('tables', [])
        if table_index >= len(tables):
            return jsonify({"error": "表格索引无效"}), 400

        table = tables[table_index]
        aligner = get_header_aligner()

        main_table = table.get('main_table', [])
        sub_tables = table.get('sub_tables', [])

        optimization_result = {
            "main_table_optimized": None,
            "sub_tables_optimized": [],
            "header_alignment": {
                "main_table": None,
                "sub_tables": []
            }
        }

        if main_table and len(main_table) > 0:
            print(f"\n[二次优化] 开始优化主表: {table_name}")
            # 现在optimize_table内部会先做硬编码去重
            opt_result = parser_system.secondary_parser.optimize_table(main_table, f"{table_name}主表")
            optimized_main = opt_result.get("optimized_table", main_table)

            aligned_main, align_stats = aligner.align_table_headers(optimized_main)
            aligned_main_with_units = aligner.align_and_merge_units(optimized_main)

            optimization_result["main_table_optimized"] = {
                "optimized": aligned_main_with_units,
                "has_optimization": opt_result.get("has_optimization", False),
                "optimization_type": opt_result.get("optimization_type", []),
                "explanation": opt_result.get("explanation", ""),
                "deduplication_stats": opt_result.get("deduplication_stats", None),  # 添加去重统计
                "header_alignment": {
                    "stats": align_stats,
                    "changes": [d for d in align_stats.get('details', []) if d['original'] != d['aligned']]
                }
            }

        for sub_idx, sub_tab in enumerate(sub_tables):
            sub_data = sub_tab.get('data', [])
            if sub_data and len(sub_data) > 0:
                print(f"\n[二次优化] 开始优化附表{sub_idx + 1}")
                opt_result = parser_system.secondary_parser.optimize_table(sub_data, f"{table_name}附表{sub_idx + 1}")
                optimized_sub = opt_result.get("optimized_table", sub_data)

                aligned_sub, sub_align_stats = aligner.align_table_headers(optimized_sub)
                aligned_sub_with_units = aligner.align_and_merge_units(optimized_sub)

                optimization_result["sub_tables_optimized"].append({
                    "sub_table_index": sub_idx,
                    "optimized": aligned_sub_with_units,
                    "has_optimization": opt_result.get("has_optimization", False),
                    "optimization_type": opt_result.get("optimization_type", []),
                    "explanation": opt_result.get("explanation", ""),
                    "deduplication_stats": opt_result.get("deduplication_stats", None),  # 添加去重统计
                    "header_alignment": {
                        "stats": sub_align_stats,
                        "changes": [d for d in sub_align_stats.get('details', []) if d['original'] != d['aligned']]
                    }
                })

        tables[table_index]["secondary_optimized"] = optimization_result

        return jsonify({
            "success": True,
            "table_index": table_index,
            "optimization_result": optimization_result
        })

    except Exception as e:
        print(f"[错误] {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/optimize_all/<task_id>', methods=['POST'])
def optimize_all_tables(task_id):
    """对所有表格进行二次优化"""
    if task_id not in tasks:
        return jsonify({"error": "任务不存在"}), 404

    task = tasks[task_id]
    if task.status != "completed" or not task.result:
        return jsonify({"error": "任务未完成"}), 400

    try:
        tables = task.result.get('tables', [])
        aligner = get_header_aligner()
        optimized_count = 0
        total_dedup_stats = {
            "total_rows_removed": 0,
            "total_cols_removed": 0,
            "tables_with_dedup": 0
        }

        for table_idx, table in enumerate(tables):
            table_name = table.get('table_id', f'表格{table_idx}')

            main_table = table.get('main_table', [])
            sub_tables = table.get('sub_tables', [])

            optimization_result = {
                "main_table_optimized": None,
                "sub_tables_optimized": [],
                "header_alignment": {
                    "main_table": None,
                    "sub_tables": []
                }
            }

            has_any_optimization = False

            # 处理主表
            if main_table and len(main_table) > 0:
                print(f"\n[批量优化] 开始优化{table_name}主表")
                # optimize_table内部会先做硬编码去重
                opt_result = parser_system.secondary_parser.optimize_table(main_table, f"{table_name}主表")
                optimized_main = opt_result.get("optimized_table", main_table)

                aligned_main, align_stats = aligner.align_table_headers(optimized_main)
                aligned_main_with_units = aligner.align_and_merge_units(optimized_main)

                # 统计去重信息
                if opt_result.get("deduplication_stats"):
                    total_dedup_stats["total_rows_removed"] += opt_result["deduplication_stats"]["rows_removed"]
                    total_dedup_stats["total_cols_removed"] += opt_result["deduplication_stats"]["cols_removed"]
                    total_dedup_stats["tables_with_dedup"] += 1

                optimization_result["main_table_optimized"] = {
                    "optimized": aligned_main_with_units,
                    "has_optimization": opt_result.get("has_optimization", False),
                    "optimization_type": opt_result.get("optimization_type", []),
                    "explanation": opt_result.get("explanation", ""),
                    "deduplication_stats": opt_result.get("deduplication_stats", None),
                    "header_alignment": {
                        "stats": align_stats,
                        "changes": [d for d in align_stats.get('details', []) if d['original'] != d['aligned']]
                    }
                }

                if opt_result.get("has_optimization"):
                    has_any_optimization = True

            # 处理所有附表
            for sub_idx, sub_tab in enumerate(sub_tables):
                sub_data = sub_tab.get('data', [])
                if sub_data and len(sub_data) > 0:
                    print(f"\n[批量优化] 开始优化{table_name}附表{sub_idx + 1}")
                    opt_result = parser_system.secondary_parser.optimize_table(sub_data,
                                                                               f"{table_name}附表{sub_idx + 1}")
                    optimized_sub = opt_result.get("optimized_table", sub_data)

                    aligned_sub, sub_align_stats = aligner.align_table_headers(optimized_sub)
                    aligned_sub_with_units = aligner.align_and_merge_units(optimized_sub)

                    # 统计去重信息
                    if opt_result.get("deduplication_stats"):
                        total_dedup_stats["total_rows_removed"] += opt_result["deduplication_stats"]["rows_removed"]
                        total_dedup_stats["total_cols_removed"] += opt_result["deduplication_stats"]["cols_removed"]
                        total_dedup_stats["tables_with_dedup"] += 1

                    optimization_result["sub_tables_optimized"].append({
                        "sub_table_index": sub_idx,
                        "optimized": aligned_sub_with_units,
                        "has_optimization": opt_result.get("has_optimization", False),
                        "optimization_type": opt_result.get("optimization_type", []),
                        "explanation": opt_result.get("explanation", ""),
                        "deduplication_stats": opt_result.get("deduplication_stats", None),
                        "header_alignment": {
                            "stats": sub_align_stats,
                            "changes": [d for d in sub_align_stats.get('details', []) if d['original'] != d['aligned']]
                        }
                    })

                    if opt_result.get("has_optimization"):
                        has_any_optimization = True

            if has_any_optimization:
                optimized_count += 1

            tables[table_idx]["secondary_optimized"] = optimization_result

        # 返回优化结果统计
        return jsonify({
            "success": True,
            "optimized_count": optimized_count,
            "total_tables": len(tables),
            "deduplication_summary": {
                "total_rows_removed": total_dedup_stats["total_rows_removed"],
                "total_cols_removed": total_dedup_stats["total_cols_removed"],
                "tables_with_deduplication": total_dedup_stats["tables_with_dedup"]
            }
        })

    except Exception as e:
        print(f"[错误] {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


def init_parser():
    global parser_system
    try:
        parser_system = GeotechnicalParserSystem()
        return True
    except Exception as e:
        print(f"初始化失败: {e}")
        return False


if __name__ == '__main__':
    print("=" * 60)
    print("岩土报告智能解析系统 - 后端服务")
    print(f"访问地址: http://{FLASK_CONFIG['HOST']}:{FLASK_CONFIG['PORT']}")
    print("=" * 60)
    with app.app_context():
        init_parser()
    app.run(
        debug=FLASK_CONFIG['DEBUG'],
        host=FLASK_CONFIG['HOST'],
        port=FLASK_CONFIG['PORT']
    )